import json
import os
from datetime import datetime
from openpyxl import load_workbook
from openpyxl.drawing.image import Image
from docx import Document

class RenYuMaterialsGenerator:
    def __init__(self):
        self.template_dir = '期初资料1'
        self.output_base_dir = 'output'
        self.courses = []
        self.metadata = {}
    
    def load_renyu_course_info(self):
        """加载任渝的课程信息"""
        file_path = os.path.join(self.template_dir, "任渝_课程信息_2026年03月21日.json")
        if not os.path.exists(file_path):
            print(f"❌ 课程信息文件不存在: {file_path}")
            return False
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            if isinstance(data, dict) and 'courses' in data:
                self.courses = data['courses']
            else:
                self.courses = data
            
            if self.courses:
                teacher_name = self.courses[0].get('teacherName', '未知教师')
                print(f"✅ 成功加载{teacher_name}的课程信息，共 {len(self.courses)} 门课程")
            else:
                print(f"✅ 成功加载课程信息，共 {len(self.courses)} 门课程")
            return True
        except Exception as e:
            print(f"❌ 加载课程信息失败: {e}")
            return False
    
    def get_course_leader(self, course):
        """根据课程信息获取课程负责人"""
        return course.get('teacherName', '未知教师')
    
    def generate_attachment3(self, course):
        """生成附件3：课程组期初教学资料检查情况记录表"""
        template_file = os.path.join(self.template_dir, "附件3 课程组教学资料检查情况记录表-完全一致模板.xlsx")
        if not os.path.exists(template_file):
            print(f"❌ 附件3模板文件不存在: {template_file}")
            return None
        
        try:
            wb = load_workbook(template_file)
            ws = wb.active
            
            course_code = course.get('courseCode', '')
            course_name = course.get('courseName', '')
            department = course.get('department', '')
            applicable_scope = course.get('applicableScope', '')
            teacher_name = course.get('teacherName', '')
            course_leader = self.get_course_leader(course)
            
            # 设置标题行（使用占位符转换方法）
            title_cell = ws.cell(row=1, column=1)
            title_cell.value = f"{course.get('formattedSemester', '2026年春季学期')}（{course_name}）课程组期初教学资料检查情况记录表"
            
            # 解析applicableScope提取归属校区和使用年级/层次/专业
            campus = ""
            grade_major = applicable_scope
            
            if applicable_scope.startswith(('东区', '绵阳校区', '德阳校区')):
                if applicable_scope.startswith('东区'):
                    campus = '东区'
                    grade_major = applicable_scope[2:]
                elif applicable_scope.startswith('绵阳校区'):
                    campus = '绵阳校区'
                    grade_major = applicable_scope[5:]
                elif applicable_scope.startswith('德阳校区'):
                    campus = '德阳校区'
                    grade_major = applicable_scope[5:]
            else:
                campus = '西区'
                grade_major = applicable_scope
            
            # 在第5行开始填充数据
            ws.cell(row=5, column=1, value='1')
            ws.cell(row=5, column=2, value=course_code)
            ws.cell(row=5, column=3, value=course_name)
            ws.cell(row=5, column=4, value=department)
            ws.cell(row=5, column=5, value=grade_major)
            ws.cell(row=5, column=6, value=campus)
            ws.cell(row=5, column=7, value=teacher_name)
            ws.cell(row=5, column=8, value=course_leader)
            
            # 生成分数，满足总分86-92之间
            import random
            while True:
                goal_score = random.randint(12, 14)
                norm_score = random.randint(8, 9)
                method_score = random.randint(8, 9)
                schedule_score = random.randint(12, 14)
                design_score = random.randint(42, 48)
                
                total = goal_score + norm_score + method_score + schedule_score + design_score
                
                if 86 <= total <= 92:
                    break
            
            ws.cell(row=5, column=9, value=goal_score)
            ws.cell(row=5, column=10, value=norm_score)
            ws.cell(row=5, column=11, value=method_score)
            ws.cell(row=5, column=12, value=schedule_score)
            ws.cell(row=5, column=13, value=design_score)
            ws.cell(row=5, column=14, value=total)
            
            # 更新资料份数
            teacher_count = 1
            
            materials_cell = ws.cell(row=19, column=5)
            materials_cell.value = f'资料 {teacher_count} 份'
            
            # 更新日期
            current_date = self.metadata.get('date', '')
            if current_date:
                date_cell = ws.cell(row=21, column=8)
                date_cell.value = f"日期：{current_date}"
            
            # 添加电子签名图片（课程组负责人签字后面）
            course_leader = self.get_course_leader(course)
            signature_image_path = os.path.join(self.template_dir, f"{course_leader}.png")
            
            if os.path.exists(signature_image_path):
                try:
                    img = Image(signature_image_path)
                    img.width = 100
                    img.height = 40
                    ws.add_image(img, 'M20')
                    print(f"✅ 电子签名已添加到 {course_name} 的附件3")
                except Exception as e:
                    print(f"⚠️ 添加电子签名失败: {e}")
            
            return wb
        except Exception as e:
            print(f"❌ 生成附件3失败: {e}")
            return None
    
    def generate_attachment1(self, course):
        """生成附件1：教学大纲审核汇总表"""
        template_file = os.path.join(self.template_dir, "附件1 教学大纲审核汇总表-完全一致模板.xlsx")
        if not os.path.exists(template_file):
            print(f"❌ 附件1模板文件不存在: {template_file}")
            return None
        
        try:
            wb = load_workbook(template_file)
            ws = wb.active
            
            course_code = course.get('courseCode', '')
            course_name = course.get('courseName', '')
            course_leader = self.get_course_leader(course)
            department = course.get('department', '')
            applicable_scope = course.get('applicableScope', '')
            teacher_name = course.get('teacherName', '')
            
            # 设置标题行（使用占位符转换方法）
            title_cell = ws.cell(row=1, column=1)
            title_cell.value = f"{course.get('formattedSemester', '2026年春季学期')}（{course_name}）教学大纲审核汇总表"
            
            # 解析applicableScope提取归属校区和使用年级/层次/专业
            campus = ""
            grade_major = applicable_scope
            
            if applicable_scope.startswith(('东区', '绵阳校区', '德阳校区')):
                if applicable_scope.startswith('东区'):
                    campus = '东区'
                    grade_major = applicable_scope[2:]
                elif applicable_scope.startswith('绵阳校区'):
                    campus = '绵阳校区'
                    grade_major = applicable_scope[5:]
                elif applicable_scope.startswith('德阳校区'):
                    campus = '德阳校区'
                    grade_major = applicable_scope[5:]
            else:
                campus = '西区'
                grade_major = applicable_scope
            
            # 在第5行开始填充数据
            ws.cell(row=5, column=1, value='1')
            ws.cell(row=5, column=2, value=course_code)
            ws.cell(row=5, column=3, value=course_name)
            ws.cell(row=5, column=4, value=department)
            ws.cell(row=5, column=5, value=campus)
            ws.cell(row=5, column=6, value=grade_major)
            ws.cell(row=5, column=7, value=teacher_name)
            ws.cell(row=5, column=8, value=course_leader)
            
            # 生成分数，满足总分86-92之间
            import random
            while True:
                goal_score = random.randint(12, 14)
                norm_score = random.randint(8, 9)
                method_score = random.randint(8, 9)
                schedule_score = random.randint(12, 14)
                design_score = random.randint(42, 48)
                
                total = goal_score + norm_score + method_score + schedule_score + design_score
                
                if 86 <= total <= 92:
                    break
            
            ws.cell(row=5, column=9, value=goal_score)
            ws.cell(row=5, column=10, value=norm_score)
            ws.cell(row=5, column=11, value=method_score)
            ws.cell(row=5, column=12, value=schedule_score)
            ws.cell(row=5, column=13, value=design_score)
            ws.cell(row=5, column=14, value=total)
            
            # 计算资料份数
            teacher_count = 1
            
            materials_cell = ws.cell(row=18, column=1)
            materials_cell.value = f'资料 {teacher_count} 份'
            
            # 更新日期
            current_date = self.metadata.get('date', '')
            if current_date:
                date_cell = ws.cell(row=20, column=10)
                date_cell.value = f"日期：{current_date}"
            
            # 添加电子签名图片（只有任渝的文件才使用陈蔚.jpg）
            teacher_name = course.get('teacherName', '')
            if teacher_name == '任渝':
                signature_image_path = os.path.join(self.template_dir, "陈蔚.jpg")
                
                if os.path.exists(signature_image_path):
                    try:
                        img = Image(signature_image_path)
                        img.width = 100
                        img.height = 40
                        ws.add_image(img, 'J19')
                        print(f"✅ 电子签名已添加到 {course_name} 的附件1")
                    except Exception as e:
                        print(f"⚠️ 添加电子签名失败: {e}")
            
            return wb
        except Exception as e:
            print(f"❌ 生成附件1失败: {e}")
            return None
    
    def generate_syllabus(self, course):
        """生成教学大纲"""
        teacher_name = course.get('teacherName', '')
        
        # 根据教师姓名选择不同的模板
        if teacher_name == '任渝':
            template_file = os.path.join(self.template_dir, "1 课程教学大纲基础模版.docx")
        else:
            template_file = os.path.join(self.template_dir, "2 课程教学大纲基础模版_副本.docx")
            
        if not os.path.exists(template_file):
            print(f"❌ 教学大纲模板文件不存在: {template_file}")
            return None
        
        try:
            doc = Document(template_file)
            
            course_code = course.get('courseCode', '')
            course_name = course.get('courseName', '')
            teacher_name = course.get('teacherName', '')
            department = course.get('department', '')
            credits = course.get('credits', '')
            total_hours = course.get('totalHours', '')
            course_nature = course.get('courseNature', '')
            applicable_scope = course.get('applicableScope', '')
            english_name = course.get('englishName', '')
            
            # 打印调试信息
            print(f"📋 模板文件: {template_file}")
            print(f"📋 教师姓名: {teacher_name}")
            print(f"📋 课程名称: {course_name}")
            print(f"📋 课程代码: {course_code}")
            print(f"📋 文档段落数量: {len(doc.paragraphs)}")
            print(f"📋 文档表格数量: {len(doc.tables)}")
            
            # 打印所有段落的内容，查找开课单位相关内容
            for i, paragraph in enumerate(doc.paragraphs):
                if '开课单位' in paragraph.text or '审批' in paragraph.text or '签名' in paragraph.text:
                    print(f"📋 段落{i}: {paragraph.text}")
            
            # 打印表格信息
            for table_idx, table in enumerate(doc.tables):
                print(f"📋 表格{table_idx}: {len(table.rows)}行")
                for row_idx, row in enumerate(table.rows):
                    print(f"📋 表格{table_idx}行{row_idx}: {len(row.cells)}列")
                    for cell_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            print(f"📋 表格{table_idx}行{row_idx}列{cell_idx}: {cell.text[:50]}...")
            
            # 替换占位符（保持字体格式）
            for paragraph in doc.paragraphs:
                # 处理{{courseName}}课程占位符（被分成两个run）
                if '{{courseName}}课程' in paragraph.text:
                    for run in paragraph.runs:
                        if '{{courseName}}' in run.text:
                            run.text = course_name
                    continue
                    
                # 处理{{formattedSemester}}学期占位符（被分成多个run）
                elif '{{formattedSemester}}学期' in paragraph.text:
                    # 找到包含占位符的run并替换，去掉重复的"学期"
                    for run in paragraph.runs:
                        if '{{' in run.text:
                            run.text = ''
                        elif 'formattedSemester' in run.text:
                            # 获取学期值（不带"学期"后缀）
                            semester = course.get('formattedSemester', '2026年春季学期')
                            if semester.endswith('学期'):
                                semester = semester[:-2]
                            run.text = semester
                        elif '}}' in run.text:
                            run.text = ''
                    continue
                    
                # 处理其他占位符
                for run in paragraph.runs:
                    if '{{courseName}}' in run.text:
                        run.text = course_name
                    elif '{{courseCode}}' in run.text:
                        run.text = course_code
                    elif '{{teacherName}}' in run.text:
                        run.text = teacher_name
                    elif '{{department}}' in run.text:
                        run.text = department
                    elif '{{credits}}' in run.text:
                        run.text = credits
                    elif '{{totalHours}}' in run.text:
                        run.text = total_hours
                    elif '{{courseNature}}' in run.text:
                        run.text = course_nature
                    elif '{{applicableScope}}' in run.text:
                        run.text = applicable_scope
                    elif '{{englishName}}' in run.text:
                        run.text = english_name
                    elif '{{formattedSemester}}' in run.text:
                        run.text = course.get('formattedSemester', '2026年春季学期')
                    elif '{{semester}}' in run.text:
                        run.text = self.metadata.get('semester', '2026年春季学期')
                    elif '{{date}}' in run.text:
                        run.text = self.metadata.get('date', '')
            
            # 处理表格中的占位符（保持字体格式）
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            
                            # 检查是否包含占位符
                            has_placeholder = False
                            new_text = text
                            
                            if '{{englishName}}' in text:
                                new_text = text.replace('{{englishName}}', english_name)
                                has_placeholder = True
                            elif '{{courseCode}}' in text:
                                new_text = text.replace('{{courseCode}}', course_code)
                                has_placeholder = True
                            elif '{{department}}' in text:
                                new_text = text.replace('{{department}}', department)
                                has_placeholder = True
                            elif '{{applicableScope}}' in text:
                                new_text = text.replace('{{applicableScope}}', applicable_scope)
                                has_placeholder = True
                            elif '{{teacherName}}' in text:
                                new_text = text.replace('{{teacherName}}', teacher_name)
                                has_placeholder = True
                            elif '{{credits}}' in text:
                                new_text = text.replace('{{credits}}', credits)
                                has_placeholder = True
                            elif '{{totalHours}}' in text:
                                new_text = text.replace('{{totalHours}}', total_hours)
                                has_placeholder = True
                            elif '{{currentDate}}' in text:
                                new_text = text.replace('{{currentDate}}', self.metadata.get('date', ''))
                                has_placeholder = True
                            elif '{{courseName}}' in text:
                                new_text = text.replace('{{courseName}}', course_name)
                                has_placeholder = True
                            elif '{{courseNature}}' in text:
                                new_text = text.replace('{{courseNature}}', course_nature)
                                has_placeholder = True
                            
                            # 只有在有占位符被替换时才更新文本
                            # 这样可以保留不包含占位符的内容，如开课单位审核人和负责人审批
                            if has_placeholder:
                                if len(paragraph.runs) == 1:
                                    run = paragraph.runs[0]
                                    run.text = new_text
                                else:
                                    paragraph.text = new_text
            
            # 删除所有图片（签名贴图），确保文档干净
            for paragraph in doc.paragraphs:
                # 删除段落中的所有图片
                for run in paragraph.runs:
                    if run._element.getchildren():
                        for child in run._element.getchildren():
                            if child.tag.endswith('graphic'):
                                run._element.remove(child)
            
            return doc
        except Exception as e:
            print(f"❌ 生成教学大纲失败: {e}")
            return None
    
    def generate_all_materials(self):
        """生成所有期初资料"""
        # 检查课程数据是否为空
        if not self.courses:
            print("❌ 课程数据为空")
            return False
        
        # 打印调试信息
        print(f"📋 self.courses类型: {type(self.courses)}")
        print(f"📋 self.courses内容: {self.courses}")
        print(f"📋 self.courses长度: {len(self.courses)}")
        
        # 合并相同课程代码的课程信息
        merged_courses = {}
        for i, course in enumerate(self.courses):
            print(f"📋 课程{i}类型: {type(course)}, 内容: {course}")
            if isinstance(course, dict):
                course_code = course.get('courseCode', '')
                if course_code not in merged_courses:
                    merged_courses[course_code] = course.copy()
                else:
                    # 只保留第一个教师的姓名，避免重复
                    pass
            else:
                print(f"❌ 课程{i}不是字典类型: {type(course)}")
        
        # 打印调试信息
        print(f"📋 合并后的课程数量: {len(merged_courses)}")
        for course_code, course in merged_courses.items():
            print(f"📋 课程代码: {course_code}, 课程类型: {type(course)}")
        
        # 设置元数据
        self.metadata = {
            'date': datetime.now().strftime('%Y年%m月%d日'),
            'semester': '2026年春季学期'
        }
        
        # 为每门课程生成资料
        for course_code, course in merged_courses.items():
            print(f"📋 当前课程: {course_code}, course类型: {type(course)}")
            print(f"📋 course内容: {course}")
            print(f"📋 metadata类型: {type(self.metadata)}")
            print(f"📋 metadata内容: {self.metadata}")
            
            course['formattedSemester'] = self.metadata['semester']
            
            # 创建课程文件夹（替换特殊字符）
            teacher_name = course.get('teacherName', '未知教师')
            course_name = course.get('courseName', '')
            safe_course_name = course_name.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
            course_folder_name = f"{course_code}_{safe_course_name}"
            # 创建教师文件夹
            teacher_folder = f"{teacher_name}_天府学院期初资料"
            output_dir = os.path.join(self.output_base_dir, teacher_folder, course_folder_name)
            os.makedirs(output_dir, exist_ok=True)
            
            # 生成教学大纲
            doc = self.generate_syllabus(course)
            if doc:
                teacher_name = course.get('teacherName', '未知教师')
                safe_filename_name = course_name.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
                output_file = os.path.join(output_dir, f"2026春-{safe_filename_name}-教学大纲-{teacher_name}.docx")
                doc.save(output_file)
                print(f"✅ 教学大纲已生成: {output_file}")
            
            # 生成附件3
            wb3 = self.generate_attachment3(course)
            if wb3:
                course_leader = self.get_course_leader(course)
                safe_filename_name = course_name.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
                output_file = os.path.join(output_dir, f"{safe_filename_name}-课程组期初教学资料检查情况记录表-{course_leader}.xlsx")
                wb3.save(output_file)
                print(f"✅ 附件3已生成: {output_file}")
            
            # 生成附件1
            wb1 = self.generate_attachment1(course)
            if wb1:
                department = course.get('department', '智能金融学院')
                course_leader = self.get_course_leader(course)
                output_file = os.path.join(output_dir, f"{department}-教学大纲审核汇总表-{course_leader}.xlsx")
                wb1.save(output_file)
                print(f"✅ 附件1已生成: {output_file}")
        
        if self.courses:
            teacher_name = self.courses[0].get('teacherName', '未知教师')
            print(f"\n🎉 {teacher_name}的期初资料生成完成！")
        else:
            print("\n🎉 期初资料生成完成！")
        return True

if __name__ == "__main__":
    generator = RenYuMaterialsGenerator()
    generator.generate_all_materials()
