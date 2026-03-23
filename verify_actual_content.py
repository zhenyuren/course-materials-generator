from docx import Document

# 读取生成的word文档，检查实际内容
file_path = "期初资料1/2026春季学期_期初资料_任渝/FIN0008B_财经公文写作/2026春-财经公文写作-教学大纲-任渝.docx"
doc = Document(file_path)

print("检查财经公文写作教学大纲的实际内容：")
print("-" * 50)

# 打印段落4和段落12的内容（这两个段落包含占位符）
print("段落 4:", doc.paragraphs[3].text)
print("段落 12:", doc.paragraphs[11].text)

print("\n所有包含占位符的段落：")
for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text
    if '{{' in text or '}}' in text:
        print(f"段落 {i+1}: {text}")
