from docx import Document

# 读取生成的文件，验证关键位置的内容和字体格式
file_path = "期初资料1/2026春季学期_期初资料_任渝/FIN0008B_财经公文写作/2026春-财经公文写作-教学大纲-任渝.docx"
doc = Document(file_path)

print("验证生成文件的内容和格式：")
print("=" * 60)

# 检查段落4（课程名称）
print("段落4（课程名称）:")
paragraph4 = doc.paragraphs[3]  # 索引从0开始
print(f"内容: {paragraph4.text}")
if paragraph4.runs:
    run = paragraph4.runs[0]
    font_size = run.font.size
    font_name = run.font.name
    print(f"字体大小: {font_size.pt if font_size else '未设置'}")
    print(f"字体名称: {font_name}")

print("\n" + "-" * 40 + "\n")

# 检查段落12（学期）
print("段落12（学期）:")
paragraph12 = doc.paragraphs[11]  # 索引从0开始
print(f"内容: {paragraph12.text}")
if paragraph12.runs:
    run = paragraph12.runs[0]
    font_size = run.font.size
    font_name = run.font.name
    print(f"字体大小: {font_size.pt if font_size else '未设置'}")
    print(f"字体名称: {font_name}")

print("\n" + "-" * 40 + "\n")

# 检查表格中的关键信息
print("表格中的关键信息:")
if doc.tables:
    table = doc.tables[0]
    # 英文名称
    english_name_cell = table.cell(0, 1)
    print(f"英文名称: {english_name_cell.text.strip()}")
    
    # 课程代码
    course_code_cell = table.cell(0, 3)
    print(f"课程代码: {course_code_cell.text.strip()}")
    
    # 开课单位
    department_cell = table.cell(1, 1)
    print(f"开课单位: {department_cell.text.strip()}")
    
    # 适用范围
    scope_cell = table.cell(2, 1)
    print(f"适用范围: {scope_cell.text.strip()}")
    
    # 制定人
    teacher_cell = table.cell(3, 1)
    print(f"制定人: {teacher_cell.text.strip()}")
    
    # 学分
    credits_cell = table.cell(4, 1)
    print(f"学分: {credits_cell.text.strip()}")
    
    # 总学时
    hours_cell = table.cell(4, 3)
    print(f"总学时: {hours_cell.text.strip()}")

print("\n✅ 验证完成！")
