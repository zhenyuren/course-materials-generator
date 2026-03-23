from docx import Document

# 读取生成的文件，检查是否还有占位符
file_path = "期初资料1/2026春季学期_期初资料_任渝/FIN0008B_财经公文写作/2026春-财经公文写作-教学大纲-任渝.docx"
doc = Document(file_path)

print("生成文件中的剩余占位符检查：")
print("=" * 60)

has_placeholders = False

# 检查段落中的占位符
print("段落中的占位符：")
for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text
    if '{{' in text or '}}' in text:
        print(f"段落 {i+1}: {text}")
        has_placeholders = True

# 检查表格中的占位符
print("\n表格中的占位符：")
table_count = 0
for table in doc.tables:
    table_count += 1
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                text = paragraph.text
                if '{{' in text or '}}' in text:
                    print(f"表格{table_count} 第{i+1}行第{j+1}列: {text}")
                    has_placeholders = True

if not has_placeholders:
    print("\n✅ 所有占位符已正确替换")
else:
    print("\n❌ 存在未替换的占位符")
