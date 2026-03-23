#!/usr/bin/env python3
# 检查生成的文档结构

from docx import Document
import os

generated_file = "output/佘脊_天府学院期初资料/DIM2501A_学科专业认知/2026春-学科专业认知-教学大纲-佘脊.docx"

if not os.path.exists(generated_file):
    print(f"生成的文档不存在: {generated_file}")
    exit(1)

# 读取生成的文档
doc = Document(generated_file)

print(f"生成文档的段落数量: {len(doc.paragraphs)}")
print(f"生成文档的表格数量: {len(doc.tables)}")

# 打印表格信息
for table_idx, table in enumerate(doc.tables):
    print(f"\n表格{table_idx}: {len(table.rows)}行 {len(table.columns)}列")
    for row_idx, row in enumerate(table.rows):
        row_text = []
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            if cell_text:
                row_text.append(f"列{cell_idx}: {cell_text[:50]}")
        if row_text:
            print(f"  行{row_idx}: {' | '.join(row_text)}")
