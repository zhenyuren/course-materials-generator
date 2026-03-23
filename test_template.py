#!/usr/bin/env python3
# 测试模板文件处理

from docx import Document
import os

template_file = "期初资料1/2 课程教学大纲基础模版_副本.docx"

if not os.path.exists(template_file):
    print(f"模板文件不存在: {template_file}")
    exit(1)

# 直接读取模板文件
doc = Document(template_file)

print(f"文档段落数量: {len(doc.paragraphs)}")
print(f"文档表格数量: {len(doc.tables)}")

# 查找包含开课单位的内容
for i, paragraph in enumerate(doc.paragraphs):
    if '开课单位' in paragraph.text or '审批' in paragraph.text or '签名' in paragraph.text:
        print(f"段落{i}: {paragraph.text}")

# 检查表格内容
for table_idx, table in enumerate(doc.tables):
    print(f"表格{table_idx}: {len(table.rows)}行")
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            if '开课单位' in cell.text or '审批' in cell.text or '签名' in cell.text:
                print(f"表格{table_idx}行{row_idx}列{cell_idx}: {cell.text}")

# 直接保存文档
output_file = "test_template_output.docx"
doc.save(output_file)
print(f"测试文档已保存: {output_file}")
