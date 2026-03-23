#!/usr/bin/env python3
# 简单测试：直接读取并保存模板文件

from docx import Document
import os

template_file = "期初资料1/2 课程教学大纲基础模版_副本.docx"
output_file = "test_simple_output.docx"

if not os.path.exists(template_file):
    print(f"模板文件不存在: {template_file}")
    exit(1)

# 直接读取模板文件
doc = Document(template_file)

print(f"文档段落数量: {len(doc.paragraphs)}")
print(f"文档表格数量: {len(doc.tables)}")

# 直接保存文档，不做任何修改
doc.save(output_file)
print(f"测试文档已保存: {output_file}")
