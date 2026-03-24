import json
import os
import re
from docx import Document
from docx.shared import Pt
from copy import deepcopy
import openpyxl

def replace_text_in_paragraph(paragraph, placeholders_dict):
    """在段落中替换占位符，同时保持原有格式"""
    # 获取段落的完整文本
    full_text = paragraph.text
    
    # 检查是否包含占位符
    has_placeholder = False
    for placeholder in placeholders_dict.keys():
        if placeholder in full_text:
            has_placeholder = True
            break
    
    if not has_placeholder:
        return 0
    
    # 保存原始runs的格式信息
    runs_info = []
    for run in paragraph.runs:
        runs_info.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name if run.font else None,
            'font_size': run.font.size if run.font else None,
            'color': run.font.color.rgb if run.font and run.font.color else None
        })
    
    # 替换文本中的占位符
    new_text = full_text
    for placeholder, value in placeholders_dict.items():
        new_text = new_text.replace(placeholder, str(value))
    
    # 清空原有runs
    for run in paragraph.runs:
        run.text = ''
    
    # 如果只有一个run，直接设置文本
    if len(paragraph.runs) == 1:
        paragraph.runs[0].text = new_text
        return 1
    
    # 如果有多个runs，尝试保持格式
    # 简单处理：将所有文本放入第一个run
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        # 清空其他runs
        for run in paragraph.runs[1:]:
            run.text = ''
    
    return 1

def replace_text_in_cell(cell, placeholders_dict):
    """在单元格中替换占位符，同时保持原有格式"""
    replacement_count = 0
    
    # 处理单元格中的段落
    for paragraph in cell.paragraphs:
        count = replace_text_in_paragraph(paragraph, placeholders_dict)
        replacement_count += count
    
    return replacement_count

import random

def fill_excel_template(course, template_file, output_path):
    try:
        print(f"  正在读取Excel模板: {template_file}")
        wb = openpyxl.load_workbook(template_file)
        
        # 构建占位符字典
        placeholders_dict = {}
        for key, value in course.items():
            placeholders_dict[f"{{{{{key}}}}}"] = str(value)  # 双花括号
            placeholders_dict[f"{{{key}}}"] = str(value)     # 单花括号
        
        # 添加默认值以处理缺失的占位字段
        default_values = {
            "campus": "东区",
            "gradeMajor": "2024级本科",
            "courseLeader": "",
            "teacherName": course.get('teacherName', ''),
            "courseCode": course.get('courseCode', ''),
            "courseName": course.get('courseName', ''),
            "department": course.get('department', ''),
            "applicableScope": course.get('applicableScope', ''),
            "credits": course.get('credits', ''),
            "totalHours": course.get('totalHours', ''),
            "courseNature": course.get('courseNature', ''),
            "currentDate": course.get('currentDate', ''),
            "englishName": course.get('englishName', ''),
            "formattedSemester": course.get('formattedSemester', '')
        }
        
        # 添加默认值到占位符字典
        for key, value in default_values.items():
            if f"{{{{{key}}}}}" not in placeholders_dict:
                placeholders_dict[f"{{{{{key}}}}}"] = str(value)
            if f"{{{key}}}" not in placeholders_dict:
                placeholders_dict[f"{{{key}}}"] = str(value)
        
        # 生成随机分数
        def generate_scores():
            # 定义各分项的满分
            max_scores = {
                'teaching_outline': 15,    # 教学大纲符合度
                'writing_standard': 10,    # 撰写规范性
                'teaching_goal': 15,       # 教学目标设计
                'teaching_method': 10,     # 教学方法设计
                'teaching_process': 50     # 教学过程设计
            }
            
            # 目标总分范围
            target_min = 86
            target_max = 92
            
            # 生成随机总分
            total_score = random.randint(target_min, target_max)
            
            # 生成各分项分数
            # 先为每个分项分配一个基础分数（满分的80%）
            base_scores = {
                'teaching_outline': int(max_scores['teaching_outline'] * 0.8),
                'writing_standard': int(max_scores['writing_standard'] * 0.8),
                'teaching_goal': int(max_scores['teaching_goal'] * 0.8),
                'teaching_method': int(max_scores['teaching_method'] * 0.8),
                'teaching_process': int(max_scores['teaching_process'] * 0.8)
            }
            
            # 计算基础总分
            base_total = sum(base_scores.values())
            
            # 计算剩余分数
            remaining = total_score - base_total
            
            # 分配剩余分数
            if remaining > 0:
                # 随机分配剩余分数，确保不超过各分项的满分
                while remaining > 0:
                    # 随机选择一个分项
                    items = list(base_scores.keys())
                    item = random.choice(items)
                    
                    # 计算该分项还能增加的分数
                    can_add = max_scores[item] - base_scores[item]
                    if can_add > 0:
                        # 随机增加1-2分
                        add = min(random.randint(1, 2), can_add, remaining)
                        base_scores[item] += add
                        remaining -= add
            
            return {
                'teaching_outline': base_scores['teaching_outline'],
                'writing_standard': base_scores['writing_standard'],
                'teaching_goal': base_scores['teaching_goal'],
                'teaching_method': base_scores['teaching_method'],
                'teaching_process': base_scores['teaching_process'],
                'total': total_score
            }
        
        # 生成分数
        scores = generate_scores()
        
        # 处理所有工作表
        for sheet in wb.worksheets:
            print(f"  处理工作表: {sheet.title}")
            
            # 遍历所有单元格
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        # 先替换所有双花括号占位符
                        for placeholder, value in placeholders_dict.items():
                            if placeholder in cell.value:
                                cell.value = cell.value.replace(placeholder, value)
                        # 处理可能存在的其他格式的占位符
                        cell.value = cell.value.replace("{{campus}}", default_values["campus"])
                        cell.value = cell.value.replace("{{gradeMajor}}", default_values["gradeMajor"])
                        cell.value = cell.value.replace("{{courseLeader}}", default_values["courseLeader"])
            
            # 填充分数（根据表格结构，假设分数在第6行）
            for row in sheet.iter_rows(min_row=5, max_row=6):
                # 假设教学大纲符合度在H列，撰写规范性在I列，教学目标设计在J列，教学方法设计在K列，教学过程设计在L列，总分在M列
                for cell in row:
                    if cell.column_letter == 'H':  # 教学大纲符合度
                        cell.value = scores['teaching_outline']
                    elif cell.column_letter == 'I':  # 撰写规范性
                        cell.value = scores['writing_standard']
                    elif cell.column_letter == 'J':  # 教学目标设计
                        cell.value = scores['teaching_goal']
                    elif cell.column_letter == 'K':  # 教学方法设计
                        cell.value = scores['teaching_method']
                    elif cell.column_letter == 'L':  # 教学过程设计
                        cell.value = scores['teaching_process']
                    elif cell.column_letter == 'M':  # 总分
                        cell.value = scores['total']
        
        print(f"  正在保存Excel文件: {output_path}")
        wb.save(output_path)
        print(f"  Excel文件保存成功!")
        return True
    except Exception as e:
        print(f"  填充Excel模板时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def fill_word_template(json_file, template_file, output_dir):
    try:
        print(f"正在读取JSON文件: {json_file}")
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        print(f"JSON数据内容: {json.dumps(data, ensure_ascii=False, indent=2)}")
        
        # 获取课程信息
        courses = data.get('courses', [])
        metadata = data.get('metadata', {})
        
        if not courses:
            print("错误: 没有找到课程信息")
            return False
        
        # 合并相同课程代码的课程信息
        merged_courses = {}
        for course in courses:
            course_code = course.get('courseCode', '')
            if course_code not in merged_courses:
                merged_courses[course_code] = course.copy()
            else:
                # 合并班级信息
                existing_scope = merged_courses[course_code].get('applicableScope', '')
                new_scope = course.get('applicableScope', '')
                
                # 提取班级信息并合并
                existing_classes = []
                new_classes = []
                
                if existing_scope:
                    existing_classes = existing_scope.replace('班级:', '').split('班级')
                if new_scope:
                    new_classes = new_scope.replace('班级:', '').split('班级')
                
                # 合并班级列表并去重
                all_classes = set(existing_classes + new_classes)
                all_classes.discard('')  # 移除空字符串
                
                # 重新生成班级信息
                merged_scope = '班级:' + '班级'.join(sorted(all_classes))
                merged_courses[course_code]['applicableScope'] = merged_scope
        
        # 转换为列表
        courses = list(merged_courses.values())
        
        print(f"原始课程记录数: {data.get('metadata', {}).get('total_courses', len(data.get('courses', [])))}")
        print(f"合并后课程数: {len(courses)}")
        for i, course in enumerate(courses):
            print(f"  {i+1}. {course.get('courseName', '')} ({course.get('courseCode', '')}) - {len(course.get('applicableScope', '').split('班级'))-1}个班级")
        
        # 定义模板文件路径
        template_base_dir = r"D:\智能金融学院\智能体设计\期初资料1"
        
        # 处理每门课程
        for i, course in enumerate(courses):
            print(f"\n正在处理第 {i + 1} 门课程...")
            print(f"课程名称: {course.get('courseName', '未知课程')}")
            
            # 根据教师姓名选择模板文件
            teacher_name = course.get('teacherName', '')
            if teacher_name == '任渝':
                word_template = os.path.join(template_base_dir, "1 课程教学大纲基础模版.docx")
                print(f"  使用模板: 1 课程教学大纲基础模版.docx (任渝专用)")
            else:
                word_template = os.path.join(template_base_dir, "2 课程教学大纲基础模版_副本.docx")
                print(f"  使用模板: 2 课程教学大纲基础模版_副本.docx (通用模板)")
            
            # 检查模板文件是否存在
            if not os.path.exists(word_template):
                print(f"  错误: Word模板文件不存在: {word_template}")
                print("  降级使用默认模板...")
                word_template = template_file
            
            # 读取Word模板
            doc = Document(word_template)
            
            # 构建占位符字典
            placeholders_dict = {}
            for key, value in course.items():
                placeholders_dict[f"{{{{{key}}}}}"] = str(value)  # 双花括号
                placeholders_dict[f"{{{key}}}"] = str(value)     # 单花括号
            
            replacement_count = 0
            
            # 替换段落中的占位符
            for paragraph in doc.paragraphs:
                count = replace_text_in_paragraph(paragraph, placeholders_dict)
                replacement_count += count
                if count > 0:
                    print(f"  段落替换完成: {paragraph.text[:50]}...")
            
            # 替换表格中的占位符
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        count = replace_text_in_cell(cell, placeholders_dict)
                        replacement_count += count
                        if count > 0:
                            print(f"  表格[{table_idx}][{row_idx}][{cell_idx}]替换完成")
            
            print(f"  总共进行了 {replacement_count} 次替换")
            
            # 获取教师相关信息
            course_code = course.get('courseCode', '')
            course_name = course.get('courseName', '未知课程')
            teacher_name = course.get('teacherName', '未知教师')
            department = course.get('department', '未知学院')
            teacher_id = course.get('teacherId', '')  # 工号
            
            # 清理文件名中的特殊字符
            safe_course_name = "".join([c for c in course_name if c.isalnum() or c in (' ', '-', '_')])
            safe_teacher_name = "".join([c for c in teacher_name if c.isalnum()])
            safe_department = "".join([c for c in department if c.isalnum() or c in (' ', '-', '_')])
            
            # 生成教师文件夹名称：学院_教师姓名_工号
            teacher_folder_name = f"{safe_department}_{safe_teacher_name}_{teacher_id}" if teacher_id else f"{safe_department}_{safe_teacher_name}"
            teacher_folder_path = os.path.join(output_dir, teacher_folder_name)
            
            # 确保教师文件夹存在
            os.makedirs(teacher_folder_path, exist_ok=True)
            print(f"  已创建/使用文件夹: {teacher_folder_name}")
            
            # 生成课程文件夹名称：课程代码_课程名
            course_folder_name = f"{course_code}_{safe_course_name}"
            course_folder_path = os.path.join(teacher_folder_path, course_folder_name)
            
            # 确保课程文件夹存在
            os.makedirs(course_folder_path, exist_ok=True)
            print(f"  已创建/使用课程文件夹: {course_folder_name}")
            
            # 生成Word大纲文件
            word_output_filename = f"{course_code}_{safe_course_name}_大纲_{safe_teacher_name}.docx"
            word_output_path = os.path.join(course_folder_path, word_output_filename)
            
            print(f"  正在保存Word文件: {word_output_path}")
            doc.save(word_output_path)
            print(f"  Word文件保存成功!")
            
            # 生成Excel文件1: 附件1 教学大纲审核汇总表
            excel1_template = os.path.join(template_base_dir, "附件1 教学大纲审核汇总表-完全一致模板.xlsx")
            excel1_output_filename = f"{course_code}_{safe_course_name}_审核汇总表_{safe_teacher_name}.xlsx"
            excel1_output_path = os.path.join(course_folder_path, excel1_output_filename)
            
            if os.path.exists(excel1_template):
                print(f"  正在生成Excel文件1: {excel1_output_filename}")
                fill_excel_template(course, excel1_template, excel1_output_path)
            else:
                print(f"  警告: Excel模板1不存在: {excel1_template}")
            
            # 生成Excel文件2: 附件3 课程组教学资料检查情况记录表
            excel2_template = os.path.join(template_base_dir, "附件3 课程组教学资料检查情况记录表-完全一致模板.xlsx")
            excel2_output_filename = f"{course_code}_{safe_course_name}_检查情况记录表_{safe_teacher_name}.xlsx"
            excel2_output_path = os.path.join(course_folder_path, excel2_output_filename)
            
            if os.path.exists(excel2_template):
                print(f"  正在生成Excel文件2: {excel2_output_filename}")
                fill_excel_template(course, excel2_template, excel2_output_path)
            else:
                print(f"  警告: Excel模板2不存在: {excel2_template}")
        
        return True
        
    except Exception as e:
        print(f"填充模板时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def process_all_json_files(base_dir):
    """处理目录中的所有JSON文件"""
    print("=" * 80)
    print("🚀 批量大纲生成器")
    print("=" * 80)
    print(f"处理目录: {base_dir}")
    print("=" * 80)
    
    # 查找所有JSON文件
    json_files = [f for f in os.listdir(base_dir) if f.endswith('.json') and f != 'import json.py']
    
    if not json_files:
        print("错误: 没有找到JSON文件")
        return
    
    print(f"找到 {len(json_files)} 个JSON文件:")
    for i, json_file in enumerate(json_files):
        print(f"  {i+1}. {json_file}")
    print("=" * 80)
    
    template_file = os.path.join(base_dir, "1 课程教学大纲基础模版.docx")
    output_dir = base_dir
    
    # 处理每个JSON文件
    for json_file in json_files:
        json_path = os.path.join(base_dir, json_file)
        print(f"\n处理文件: {json_file}")
        print("-" * 80)
        success = fill_word_template(json_path, template_file, output_dir)
        if success:
            print(f"✅ {json_file} 处理成功")
        else:
            print(f"❌ {json_file} 处理失败")
        print("-" * 80)
    
    print("=" * 80)
    print(f"批量处理完成！共处理了 {len(json_files)} 个JSON文件")
    print("=" * 80)

def upload_files_to_nuwa(base_dir):
    """上传生成的大纲文件到女娲平台"""
    try:
        # 导入上传模块
        from upload_to_nuwa import upload_all_outlines
        
        # 使用实际的API Key
        api_key = "ak-e7cc7011dce24aa88ed0ed2bba49c90f"
        
        # 上传文件
        upload_all_outlines(base_dir, api_key)
        
    except ImportError:
        print("⚠️  上传模块未找到，请确保 upload_to_nuwa.py 文件存在")
    except Exception as e:
        print(f"⚠️  上传文件时出错: {str(e)}")

def main():
    # 定义路径
    base_dir = r"D:\智能金融学院\智能体设计"
    process_all_json_files(base_dir)
    
    # 上传文件到女娲平台
    print("\n" + "=" * 80)
    print("📤 准备上传文件到女娲平台")
    print("=" * 80)
    upload_files_to_nuwa(base_dir)

if __name__ == "__main__":
    main()