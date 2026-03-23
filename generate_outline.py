import json
import os
import re
from docx import Document
from docx.shared import Pt
from copy import deepcopy

def replace_text_in_paragraph(paragraph, placeholders_dict):
    """在段落中替换占位符，同时保持原有格式"""
    # 获取段落的完整文本
    full_text = paragraph.text
    
    # 检查是否包含占位符
    has_placeholder = False
    for placeholder in placeholders_dict.keys():
        if placeholder in full_text:
            has_placeholder = True
            break
    
    if not has_placeholder:
        return 0
    
    # 保存原始runs的格式信息
    runs_info = []
    for run in paragraph.runs:
        runs_info.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name if run.font else None,
            'font_size': run.font.size if run.font else None,
            'color': run.font.color.rgb if run.font and run.font.color else None
        })
    
    # 替换文本中的占位符
    new_text = full_text
    for placeholder, value in placeholders_dict.items():
        new_text = new_text.replace(placeholder, str(value))
    
    # 清空原有runs
    for run in paragraph.runs:
        run.text = ''
    
    # 如果只有一个run，直接设置文本
    if len(paragraph.runs) == 1:
        paragraph.runs[0].text = new_text
        return 1
    
    # 如果有多个runs，尝试保持格式
    # 简单处理：将所有文本放入第一个run
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        # 清空其他runs
        for run in paragraph.runs[1:]:
            run.text = ''
    
    return 1

def replace_text_in_cell(cell, placeholders_dict):
    """在单元格中替换占位符，同时保持原有格式"""
    replacement_count = 0
    
    # 处理单元格中的段落
    for paragraph in cell.paragraphs:
        count = replace_text_in_paragraph(paragraph, placeholders_dict)
        replacement_count += count
    
    return replacement_count

def fill_word_template(json_file, template_file, output_dir):
    try:
        print(f"正在读取JSON文件: {json_file}")
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        print(f"JSON数据内容: {json.dumps(data, ensure_ascii=False, indent=2)}")
        
        # 获取课程信息
        courses = data.get('courses', [])
        metadata = data.get('metadata', {})
        
        if not courses:
            print("错误: 没有找到课程信息")
            return False
        
        # 合并相同课程代码的课程信息
        merged_courses = {}
        for course in courses:
            course_code = course.get('courseCode', '')
            if course_code not in merged_courses:
                merged_courses[course_code] = course.copy()
            else:
                # 合并班级信息
                existing_scope = merged_courses[course_code].get('applicableScope', '')
                new_scope = course.get('applicableScope', '')
                
                # 提取班级信息并合并
                existing_classes = []
                new_classes = []
                
                if existing_scope:
                    existing_classes = existing_scope.replace('班级:', '').split('班级')
                if new_scope:
                    new_classes = new_scope.replace('班级:', '').split('班级')
                
                # 合并班级列表并去重
                all_classes = set(existing_classes + new_classes)
                all_classes.discard('')  # 移除空字符串
                
                # 重新生成班级信息
                merged_scope = '班级:' + '班级'.join(sorted(all_classes))
                merged_courses[course_code]['applicableScope'] = merged_scope
        
        # 转换为列表
        courses = list(merged_courses.values())
        
        print(f"原始课程记录数: {data.get('metadata', {}).get('total_courses', len(data.get('courses', [])))}")
        print(f"合并后课程数: {len(courses)}")
        for i, course in enumerate(courses):
            print(f"  {i+1}. {course.get('courseName', '')} ({course.get('courseCode', '')}) - {len(course.get('applicableScope', '').split('班级'))-1}个班级")
        
        # 检查模板文件是否存在
        if not os.path.exists(template_file):
            print(f"错误: 模板文件不存在: {template_file}")
            return False
        
        print(f"正在读取模板文件: {template_file}")
        
        # 处理每门课程
        for i, course in enumerate(courses):
            print(f"\n正在处理第 {i + 1} 门课程...")
            print(f"课程名称: {course.get('courseName', '未知课程')}")
            
            # 根据教师姓名选择模板文件
            teacher_name = course.get('teacherName', '')
            if teacher_name == '任渝':
                current_template = os.path.join(output_dir, "1 课程教学大纲基础模版.docx")
                print(f"  使用模板: 1 课程教学大纲基础模版.docx (任渝专用)")
            else:
                current_template = os.path.join(output_dir, "1 课程教学大纲基础模版2.docx")
                print(f"  使用模板: 1 课程教学大纲基础模版2.docx (通用模板)")
            
            # 检查模板文件是否存在
            if not os.path.exists(current_template):
                print(f"  错误: 模板文件不存在: {current_template}")
                print("  降级使用默认模板...")
                current_template = template_file
            
            doc = Document(current_template)
            
            # 构建占位符字典
            placeholders_dict = {}
            for key, value in course.items():
                placeholders_dict[f"{{{{{key}}}}}"] = str(value)  # 双花括号
                placeholders_dict[f"{{{key}}}"] = str(value)     # 单花括号
            
            replacement_count = 0
            
            # 替换段落中的占位符
            for paragraph in doc.paragraphs:
                count = replace_text_in_paragraph(paragraph, placeholders_dict)
                replacement_count += count
                if count > 0:
                    print(f"  段落替换完成: {paragraph.text[:50]}...")
            
            # 替换表格中的占位符
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        count = replace_text_in_cell(cell, placeholders_dict)
                        replacement_count += count
                        if count > 0:
                            print(f"  表格[{table_idx}][{row_idx}][{cell_idx}]替换完成")
            
            print(f"  总共进行了 {replacement_count} 次替换")
            
            # 获取教师相关信息
            course_code = course.get('courseCode', '')
            course_name = course.get('courseName', '未知课程')
            teacher_name = course.get('teacherName', '未知教师')
            department = course.get('department', '未知学院')
            teacher_id = course.get('teacherId', '')  # 工号
            
            # 清理文件名中的特殊字符
            safe_course_name = "".join([c for c in course_name if c.isalnum() or c in (' ', '-', '_')])
            safe_teacher_name = "".join([c for c in teacher_name if c.isalnum()])
            safe_department = "".join([c for c in department if c.isalnum() or c in (' ', '-', '_')])
            
            # 生成教师文件夹名称：学院_教师姓名_工号
            teacher_folder_name = f"{safe_department}_{safe_teacher_name}_{teacher_id}" if teacher_id else f"{safe_department}_{safe_teacher_name}"
            teacher_folder_path = os.path.join(output_dir, teacher_folder_name)
            
            # 确保教师文件夹存在
            os.makedirs(teacher_folder_path, exist_ok=True)
            print(f"  已创建/使用文件夹: {teacher_folder_name}")
            
            # 生成输出文件名：课程代码_课程名称_大纲_教师姓名
            output_filename = f"{course_code}_{safe_course_name}_大纲_{safe_teacher_name}.docx"
            output_path = os.path.join(teacher_folder_path, output_filename)
            
            print(f"  正在保存文件: {output_path}")
            doc.save(output_path)
            print(f"  文件保存成功!")
        
        return True
        
    except Exception as e:
        print(f"填充模板时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def main():
    # 定义路径
    base_dir = r"D:\智能金融学院\智能体设计"
    
    # 自动查找最新的JSON文件
    json_files = [f for f in os.listdir(base_dir) if f.endswith('.json') and f != 'import json.py']
    if not json_files:
        print("错误: 没有找到JSON文件")
        return
    
    # 按修改时间排序，取最新的
    json_files.sort(key=lambda x: os.path.getmtime(os.path.join(base_dir, x)), reverse=True)
    latest_json = os.path.join(base_dir, json_files[0])
    
    json_file = latest_json
    template_file = os.path.join(base_dir, "1 课程教学大纲基础模版.docx")
    output_dir = base_dir
    
    print("=" * 80)
    print("开始执行大纲模板生成流程")
    print("=" * 80)
    
    # 检查文件是否存在
    if not os.path.exists(json_file):
        print(f"错误: JSON文件不存在: {json_file}")
        return
    
    if not os.path.exists(template_file):
        print(f"错误: 模板文件不存在: {template_file}")
        return
    
    print(f"JSON文件: {json_file}")
    print(f"模板文件: {template_file}")
    print(f"输出目录: {output_dir}")
    print("=" * 80)
    
    # 执行模板填充
    success = fill_word_template(json_file, template_file, output_dir)
    
    print("=" * 80)
    if success:
        print("大纲模板生成完成！")
    else:
        print("大纲模板生成失败，请检查错误信息。")
    print("=" * 80)

if __name__ == "__main__":
    main()