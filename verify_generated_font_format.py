from docx import Document

# 读取生成的文件，验证字体格式
file_path = "期初资料1/2026春季学期_期初资料_任渝/FIN0008B_财经公文写作/2026春-财经公文写作-教学大纲-任渝.docx"
doc = Document(file_path)

print("生成文件中字体格式验证：")
print("=" * 60)

# 检查段落4（课程名称）
print("段落4（课程名称）:")
paragraph4 = doc.paragraphs[3]  # 索引从0开始
print(f"内容: {paragraph4.text}")
if paragraph4.runs:
    for i, run in enumerate(paragraph4.runs):
        font_size = run.font.size
        font_name = run.font.name
        bold = run.bold
        italic = run.italic
        print(f"Run {i+1}:")
        print(f"  字体大小: {font_size.pt if font_size else '未设置'}")
        print(f"  字体名称: {font_name}")
        print(f"  粗体: {bold}")
        print(f"  斜体: {italic}")

print("\n" + "-" * 40 + "\n")

# 检查段落12（学期）
print("段落12（学期）:")
paragraph12 = doc.paragraphs[11]  # 索引从0开始
print(f"内容: {paragraph12.text}")
if paragraph12.runs:
    for i, run in enumerate(paragraph12.runs):
        font_size = run.font.size
        font_name = run.font.name
        bold = run.bold
        italic = run.italic
        print(f"Run {i+1}:")
        print(f"  字体大小: {font_size.pt if font_size else '未设置'}")
        print(f"  字体名称: {font_name}")
        print(f"  粗体: {bold}")
        print(f"  斜体: {italic}")
