import json
from docx import Document
import os

def fill_word_template(json_file, template_file, output_dir):
    # 读取JSON数据
    with open(json_file, 'r', encoding='utf-8') as f:
        courses = json.load(f)
    
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)
    
    # 处理每个课程
    for i, course in enumerate(courses):
        # 打开模板文档
        doc = Document(template_file)
        
        # 替换文档中的占位符
        for paragraph in doc.paragraphs:
            for key, value in course.items():
                if f"{{{key}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in course.items():
                        if f"{{{key}}}" in cell.text:
                            cell.text = cell.text.replace(f"{{{key}}}", str(value))
        
        # 保存生成的文档
        output_file = os.path.join(output_dir, f"课程信息_{course.get('courseName', f'课程{i+1}')}.docx")
        doc.save(output_file)
        print(f"已生成: {output_file}")

if __name__ == "__main__":
    # 配置文件路径
    json_file = "course_data.json"
    template_file = "template.docx"
    output_dir = "output"
    
    fill_word_template(json_file, template_file, output_dir)