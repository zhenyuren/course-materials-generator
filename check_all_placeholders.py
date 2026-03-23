from docx import Document

# 读取模板文件，检查所有占位符
template_file = "期初资料1/1 课程教学大纲基础模版.docx"
doc = Document(template_file)

print("模板文件中的所有占位符：")
print("=" * 60)

all_placeholders = []

# 检查段落中的占位符
print("段落中的占位符：")
for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text
    if '{{' in text or '}}' in text:
        print(f"段落 {i+1}: {text}")
        parts = text.split('{{')
        for part in parts[1:]:
            if '}}' in part:
                placeholder = '{{' + part.split('}}')[0] + '}}'
                all_placeholders.append(placeholder)

# 检查表格中的占位符
print("\n表格中的占位符：")
table_count = 0
for table in doc.tables:
    table_count += 1
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                text = paragraph.text
                if '{{' in text or '}}' in text:
                    print(f"表格{table_count} 第{i+1}行第{j+1}列: {text}")
                    parts = text.split('{{')
                    for part in parts[1:]:
                        if '}}' in part:
                            placeholder = '{{' + part.split('}}')[0] + '}}'
                            all_placeholders.append(placeholder)

# 去重并显示所有占位符
print("\n所有唯一的占位符：")
unique_placeholders = list(set(all_placeholders))
for placeholder in unique_placeholders:
    print(f"- {placeholder}")
