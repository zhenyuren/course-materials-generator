from docx import Document

# 读取生成的word文档，检查表格中的占位符替换情况
file_path = "期初资料1/2026春季学期_期初资料_任渝/FIN0008B_财经公文写作/2026春-财经公文写作-教学大纲-任渝.docx"
doc = Document(file_path)

print("检查教学大纲表格中的占位符替换情况：")
print("-" * 50)

has_placeholders = False
for table in doc.tables:
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                text = paragraph.text
                if '{{' in text or '}}' in text:
                    print(f"表格第{i+1}行第{j+1}列包含未替换的占位符: {text}")
                    has_placeholders = True
                else:
                    print(f"表格第{i+1}行第{j+1}列: {text}")

if not has_placeholders:
    print("\n✅ 表格中所有占位符已正确替换")
else:
    print("\n❌ 表格中存在未替换的占位符")
