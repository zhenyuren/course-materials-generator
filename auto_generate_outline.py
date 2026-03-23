import json
import os
import time
from docx import Document
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

class OutlineGeneratorHandler(FileSystemEventHandler):
    def __init__(self, template_file, output_dir):
        self.template_file = template_file
        self.output_dir = output_dir
        self.processed_files = set()
        
    def on_created(self, event):
        if event.is_directory:
            return
            
        if event.src_path.endswith('.json'):
            print(f"\n检测到新的JSON文件: {event.src_path}")
            time.sleep(1)  # 等待文件写入完成
            self.generate_outline(event.src_path)
    
    def on_modified(self, event):
        if event.is_directory:
            return
            
        if event.src_path.endswith('.json'):
            file_path = event.src_path
            if file_path not in self.processed_files:
                print(f"\n检测到JSON文件修改: {file_path}")
                time.sleep(1)  # 等待文件写入完成
                self.generate_outline(file_path)
                self.processed_files.add(file_path)
    
    def replace_text_in_paragraph(self, paragraph, placeholders_dict):
        """在段落中替换占位符，同时保持原有格式"""
        # 获取段落的完整文本
        full_text = paragraph.text
        
        # 检查是否包含占位符
        has_placeholder = False
        for placeholder in placeholders_dict.keys():
            if placeholder in full_text:
                has_placeholder = True
                break
        
        if not has_placeholder:
            return 0
        
        # 保存原始runs的格式信息
        runs_info = []
        for run in paragraph.runs:
            runs_info.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name if run.font else None,
                'font_size': run.font.size if run.font else None,
                'color': run.font.color.rgb if run.font and run.font.color else None
            })
        
        # 替换文本中的占位符
        new_text = full_text
        for placeholder, value in placeholders_dict.items():
            new_text = new_text.replace(placeholder, str(value))
        
        # 清空原有runs
        for run in paragraph.runs:
            run.text = ''
        
        # 如果只有一个run，直接设置文本
        if len(paragraph.runs) == 1:
            paragraph.runs[0].text = new_text
            return 1
        
        # 如果有多个runs，尝试保持格式
        # 简单处理：将所有文本放入第一个run
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            # 清空其他runs
            for run in paragraph.runs[1:]:
                run.text = ''
        
        return 1
    
    def replace_text_in_cell(self, cell, placeholders_dict):
        """在单元格中替换占位符，同时保持原有格式"""
        replacement_count = 0
        
        # 处理单元格中的段落
        for paragraph in cell.paragraphs:
            count = self.replace_text_in_paragraph(paragraph, placeholders_dict)
            replacement_count += count
        
        return replacement_count
    
    def generate_outline(self, json_file):
        try:
            print(f"正在处理文件: {json_file}")
            
            # 读取JSON数据
            with open(json_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            print(f"JSON数据内容: {json.dumps(data, ensure_ascii=False, indent=2)}")
            
            # 获取课程信息
            courses = data.get('courses', [])
            metadata = data.get('metadata', {})
            
            if not courses:
                print("错误: 没有找到课程信息")
                return
            
            # 合并相同课程代码的课程信息
            merged_courses = {}
            for course in courses:
                course_code = course.get('courseCode', '')
                if course_code not in merged_courses:
                    merged_courses[course_code] = course.copy()
                else:
                    # 合并班级信息
                    existing_scope = merged_courses[course_code].get('applicableScope', '')
                    new_scope = course.get('applicableScope', '')
                    
                    # 提取班级信息并合并
                    existing_classes = []
                    new_classes = []
                    
                    if existing_scope:
                        existing_classes = existing_scope.replace('班级:', '').split('班级')
                    if new_scope:
                        new_classes = new_scope.replace('班级:', '').split('班级')
                    
                    # 合并班级列表并去重
                    all_classes = set(existing_classes + new_classes)
                    all_classes.discard('')  # 移除空字符串
                    
                    # 重新生成班级信息
                    merged_scope = '班级:' + '班级'.join(sorted(all_classes))
                    merged_courses[course_code]['applicableScope'] = merged_scope
            
            # 转换为列表
            courses = list(merged_courses.values())
            
            print(f"原始课程记录数: {data.get('metadata', {}).get('total_courses', len(data.get('courses', [])))}")
            print(f"合并后课程数: {len(courses)}")
            for i, course in enumerate(courses):
                print(f"  {i+1}. {course.get('courseName', '')} ({course.get('courseCode', '')}) - {len(course.get('applicableScope', '').split('班级'))-1}个班级")
            
            # 处理每门课程
            for i, course in enumerate(courses):
                print(f"\n正在处理第 {i + 1} 门课程...")
                print(f"课程名称: {course.get('courseName', '未知课程')}")
                
                # 根据教师姓名选择模板文件
                teacher_name = course.get('teacherName', '')
                if teacher_name == '任渝':
                    current_template = os.path.join(self.output_dir, "1 课程教学大纲基础模版.docx")
                    print(f"  使用模板: 1 课程教学大纲基础模版.docx (任渝专用)")
                else:
                    current_template = os.path.join(self.output_dir, "1 课程教学大纲基础模版2.docx")
                    print(f"  使用模板: 1 课程教学大纲基础模版2.docx (通用模板)")
                
                # 检查模板文件是否存在
                if not os.path.exists(current_template):
                    print(f"  错误: 模板文件不存在: {current_template}")
                    print("  降级使用默认模板...")
                    current_template = self.template_file
                
                doc = Document(current_template)
                
                # 构建占位符字典
                placeholders_dict = {}
                for key, value in course.items():
                    placeholders_dict[f"{{{{{key}}}}}"] = str(value)  # 双花括号
                    placeholders_dict[f"{{{key}}}"] = str(value)     # 单花括号
                
                replacement_count = 0
                
                # 替换段落中的占位符
                for paragraph in doc.paragraphs:
                    count = self.replace_text_in_paragraph(paragraph, placeholders_dict)
                    replacement_count += count
                    if count > 0:
                        print(f"  段落替换完成: {paragraph.text[:50]}...")
                
                # 替换表格中的占位符
                for table_idx, table in enumerate(doc.tables):
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            count = self.replace_text_in_cell(cell, placeholders_dict)
                            replacement_count += count
                            if count > 0:
                                print(f"  表格[{table_idx}][{row_idx}][{cell_idx}]替换完成")
                
                print(f"  总共进行了 {replacement_count} 次替换")
                
                # 获取教师相关信息
                course_code = course.get('courseCode', '')
                course_name = course.get('courseName', '未知课程')
                teacher_name = course.get('teacherName', '未知教师')
                department = course.get('department', '未知学院')
                teacher_id = course.get('teacherId', '')  # 工号
                
                # 清理文件名中的特殊字符
                safe_course_name = "".join([c for c in course_name if c.isalnum() or c in (' ', '-', '_')])
                safe_teacher_name = "".join([c for c in teacher_name if c.isalnum()])
                safe_department = "".join([c for c in department if c.isalnum() or c in (' ', '-', '_')])
                
                # 生成教师文件夹名称：学院_教师姓名_工号
                teacher_folder_name = f"{safe_department}_{safe_teacher_name}_{teacher_id}" if teacher_id else f"{safe_department}_{safe_teacher_name}"
                teacher_folder_path = os.path.join(self.output_dir, teacher_folder_name)
                
                # 确保教师文件夹存在
                os.makedirs(teacher_folder_path, exist_ok=True)
                print(f"  已创建/使用文件夹: {teacher_folder_name}")
                
                # 生成输出文件名：课程代码_课程名称_大纲_教师姓名
                output_filename = f"{course_code}_{safe_course_name}_大纲_{safe_teacher_name}.docx"
                output_path = os.path.join(teacher_folder_path, output_filename)
                
                print(f"  正在保存文件: {output_path}")
                doc.save(output_path)
                print(f"  文件保存成功!")
            
            print(f"✅ 成功生成 {len(courses)} 个大纲文件\n")
            
        except Exception as e:
            print(f"❌ 处理文件时出错: {str(e)}")
            import traceback
            traceback.print_exc()

def main():
    # 定义路径
    base_dir = r"D:\智能金融学院\智能体设计"
    template_file = os.path.join(base_dir, "1 课程教学大纲基础模版.docx")
    output_dir = base_dir
    
    print("=" * 80)
    print("🚀 大纲自动生成器已启动")
    print("=" * 80)
    print(f"监控目录: {base_dir}")
    print(f"模板文件: {template_file}")
    print(f"输出目录: {output_dir}")
    print("=" * 80)
    print("💡 提示: 将JSON文件放入此目录，将自动生成对应的大纲文件")
    print("💡 提示: 按 Ctrl+C 停止监控")
    print("=" * 80)
    
    # 检查模板文件是否存在
    if not os.path.exists(template_file):
        print(f"❌ 错误: 模板文件不存在: {template_file}")
        return
    
    # 创建事件处理器
    event_handler = OutlineGeneratorHandler(template_file, output_dir)
    observer = Observer()
    observer.schedule(event_handler, base_dir, recursive=False)
    
    # 启动监控
    observer.start()
    
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        print("\n\n⏹️  停止监控...")
        observer.stop()
    
    observer.join()
    print("✅ 监控已停止")

if __name__ == "__main__":
    main()