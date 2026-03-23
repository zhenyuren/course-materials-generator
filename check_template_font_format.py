from docx import Document

# 读取模板文件，检查占位符的字体格式
template_file = "期初资料1/1 课程教学大纲基础模版.docx"
doc = Document(template_file)

print("模板文件中占位符的字体格式：")
print("=" * 60)

# 检查段落4（{{courseName}}课程）
print("段落4（{{courseName}}课程）:")
paragraph4 = doc.paragraphs[3]  # 索引从0开始
print(f"内容: {paragraph4.text}")
if paragraph4.runs:
    for i, run in enumerate(paragraph4.runs):
        font_size = run.font.size
        font_name = run.font.name
        bold = run.bold
        italic = run.italic
        print(f"Run {i+1}:")
        print(f"  字体大小: {font_size.pt if font_size else '未设置'}")
        print(f"  字体名称: {font_name}")
        print(f"  粗体: {bold}")
        print(f"  斜体: {italic}")

print("\n" + "-" * 40 + "\n")

# 检查段落12（{{formattedSemester}}学期）
print("段落12（{{formattedSemester}}学期）:")
paragraph12 = doc.paragraphs[11]  # 索引从0开始
print(f"内容: {paragraph12.text}")
if paragraph12.runs:
    for i, run in enumerate(paragraph12.runs):
        font_size = run.font.size
        font_name = run.font.name
        bold = run.bold
        italic = run.italic
        print(f"Run {i+1}:")
        print(f"  字体大小: {font_size.pt if font_size else '未设置'}")
        print(f"  字体名称: {font_name}")
        print(f"  粗体: {bold}")
        print(f"  斜体: {italic}")
