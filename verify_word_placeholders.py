from docx import Document

# 读取生成的word文档，检查占位符是否被替换
file_path = "期初资料1/2026春季学期_期初资料_任渝/FIN0008B_财经公文写作/2026春-财经公文写作-教学大纲-任渝.docx"
doc = Document(file_path)

print("检查财经公文写作教学大纲中的占位符替换情况：")
print("-" * 50)

has_placeholders = False
for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text
    if '{{' in text or '}}' in text:
        print(f"段落 {i+1} 包含未替换的占位符: {text}")
        has_placeholders = True

if not has_placeholders:
    print("✅ 所有占位符已正确替换")
else:
    print("❌ 存在未替换的占位符")

# 打印前5个段落内容，验证替换效果
print("\n前5个段落内容：")
for i, paragraph in enumerate(doc.paragraphs[:5]):
    print(f"段落 {i+1}: {paragraph.text}")
