#!/usr/bin/env python3
# 测试：不删除图片，看看表格内容是否保留

from docx import Document
import os

template_file = "期初资料1/2 课程教学大纲基础模版_副本.docx"
output_file = "test_no_image_delete.docx"

if not os.path.exists(template_file):
    print(f"模板文件不存在: {template_file}")
    exit(1)

# 读取模板文件
doc = Document(template_file)

print(f"文档段落数量: {len(doc.paragraphs)}")
print(f"文档表格数量: {len(doc.tables)}")

# 只替换占位符，不删除图片
course_name = "测试课程"
course_code = "TEST001"
teacher_name = "测试老师"
department = "测试学院"
credits = "3.0"
total_hours = "48"
course_nature = "专业必修课"
applicable_scope = "测试专业"
english_name = "Test Course"
current_date = "2026年03月23日"

# 处理段落中的占位符
for paragraph in doc.paragraphs:
    if '{{courseName}}课程' in paragraph.text:
        for run in paragraph.runs:
            if '{{courseName}}' in run.text:
                run.text = course_name
        continue
    elif '{{formattedSemester}}学期' in paragraph.text:
        for run in paragraph.runs:
            if '{{' in run.text:
                run.text = ''
            elif 'formattedSemester' in run.text:
                run.text = '2026年春季'
            elif '}}' in run.text:
                run.text = ''
        continue
    for run in paragraph.runs:
        if '{{courseName}}' in run.text:
            run.text = course_name
        elif '{{courseCode}}' in run.text:
            run.text = course_code
        elif '{{teacherName}}' in run.text:
            run.text = teacher_name
        elif '{{department}}' in run.text:
            run.text = department
        elif '{{credits}}' in run.text:
            run.text = credits
        elif '{{totalHours}}' in run.text:
            run.text = total_hours
        elif '{{courseNature}}' in run.text:
            run.text = course_nature
        elif '{{applicableScope}}' in run.text:
            run.text = applicable_scope
        elif '{{englishName}}' in run.text:
            run.text = english_name
        elif '{{formattedSemester}}' in run.text:
            run.text = '2026年春季学期'
        elif '{{semester}}' in run.text:
            run.text = '2026年春季学期'
        elif '{{date}}' in run.text:
            run.text = current_date

# 处理表格中的占位符
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                text = paragraph.text
                has_placeholder = False
                new_text = text
                if '{{englishName}}' in text:
                    new_text = text.replace('{{englishName}}', english_name)
                    has_placeholder = True
                elif '{{courseCode}}' in text:
                    new_text = text.replace('{{courseCode}}', course_code)
                    has_placeholder = True
                elif '{{department}}' in text:
                    new_text = text.replace('{{department}}', department)
                    has_placeholder = True
                elif '{{applicableScope}}' in text:
                    new_text = text.replace('{{applicableScope}}', applicable_scope)
                    has_placeholder = True
                elif '{{teacherName}}' in text:
                    new_text = text.replace('{{teacherName}}', teacher_name)
                    has_placeholder = True
                elif '{{credits}}' in text:
                    new_text = text.replace('{{credits}}', credits)
                    has_placeholder = True
                elif '{{totalHours}}' in text:
                    new_text = text.replace('{{totalHours}}', total_hours)
                    has_placeholder = True
                elif '{{currentDate}}' in text:
                    new_text = text.replace('{{currentDate}}', current_date)
                    has_placeholder = True
                elif '{{courseName}}' in text:
                    new_text = text.replace('{{courseName}}', course_name)
                    has_placeholder = True
                elif '{{courseNature}}' in text:
                    new_text = text.replace('{{courseNature}}', course_nature)
                    has_placeholder = True
                if has_placeholder:
                    if len(paragraph.runs) == 1:
                        run = paragraph.runs[0]
                        run.text = new_text
                    else:
                        paragraph.text = new_text

# 保存文档，不删除图片
doc.save(output_file)
print(f"测试文档已保存: {output_file}")
