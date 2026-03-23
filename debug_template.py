#!/usr/bin/env python3
# 详细调试模板文件

from docx import Document
import os

template_file = "期初资料1/2 课程教学大纲基础模版_副本.docx"

if not os.path.exists(template_file):
    print(f"模板文件不存在: {template_file}")
    exit(1)

# 读取模板文件
doc = Document(template_file)

print("="*50)
print("文档结构详细分析")
print("="*50)

# 打印所有段落
print("\n【段落内容】")
for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if text:
        print(f"段落{i}: {text}")

# 打印所有表格
print("\n【表格内容】")
for table_idx, table in enumerate(doc.tables):
    print(f"\n表格{table_idx}: {len(table.rows)}行 {len(table.columns)}列")
    for row_idx, row in enumerate(table.rows):
        row_text = []
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            if cell_text:
                row_text.append(f"列{cell_idx}: {cell_text[:50]}")
        if row_text:
            print(f"  行{row_idx}: {' | '.join(row_text)}")
