from docx import Document

# 读取模板文件，检查占位符的实际内容
template_file = "期初资料1/1 课程教学大纲基础模版.docx"
doc = Document(template_file)

print("模板文件中占位符的详细信息：")
print("=" * 60)

# 检查段落4（课程名称）
print("段落4（课程名称）:")
paragraph4 = doc.paragraphs[3]  # 索引从0开始
print(f"完整内容: '{paragraph4.text}'")
print(f"内容长度: {len(paragraph4.text)}")
print(f"是否包含 '{{courseName}}课程': {'{{courseName}}课程' in paragraph4.text}")

if paragraph4.runs:
    for i, run in enumerate(paragraph4.runs):
        print(f"\nRun {i+1}:")
        print(f"  内容: '{run.text}'")
        print(f"  内容长度: {len(run.text)}")
        print(f"  是否包含 '{{courseName}}': {'{{courseName}}' in run.text}")
        print(f"  是否包含 '课程': {'课程' in run.text}")

print("\n" + "-" * 40 + "\n")

# 检查段落12（学期）
print("段落12（学期）:")
paragraph12 = doc.paragraphs[11]  # 索引从0开始
print(f"完整内容: '{paragraph12.text}'")
print(f"内容长度: {len(paragraph12.text)}")
print(f"是否包含 '{{formattedSemester}}学期': {'{{formattedSemester}}学期' in paragraph12.text}")

if paragraph12.runs:
    for i, run in enumerate(paragraph12.runs):
        print(f"\nRun {i+1}:")
        print(f"  内容: '{run.text}'")
        print(f"  内容长度: {len(run.text)}")
        print(f"  是否包含 '{{formattedSemester}}': {'{{formattedSemester}}' in run.text}")
        print(f"  是否包含 '学期': {'学期' in run.text}")
